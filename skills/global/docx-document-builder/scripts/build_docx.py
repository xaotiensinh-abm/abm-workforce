# -*- coding: utf-8 -*-
"""
ABM-Workforce HDSD -> Premium DOCX v4.1
Static TOC (pre-populated) + all P0+P1 features.
"""
import re
import sys
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml

# ===== THEME CONFIG =====
THEME = {
    'primary':         RGBColor(0x1B, 0x2A, 0x4A),
    'secondary':       RGBColor(0x2C, 0x3E, 0x6B),
    'accent':          RGBColor(0x34, 0x98, 0xDB),
    'text':            RGBColor(0x33, 0x33, 0x33),
    'muted':           RGBColor(0x88, 0x88, 0x88),
    'white':           RGBColor(0xFF, 0xFF, 0xFF),
    'bg_code':         'F5F5F5',
    'bg_table_header': '1B2A4A',
    'bg_table_alt':    'F0F4FA',
    'border_light':    'CCCCCC',
    'font_main':       'Arial',
    'font_code':       'Consolas',
    'font_size':       Pt(11),
    'code_size':       Pt(9),
    'h1_size':         Pt(22),
    'h2_size':         Pt(16),
    'h3_size':         Pt(13),
}

DOC_TITLE = 'ABM-WORKFORCE v3.0'
DOC_AUTHOR = 'ABM-Workforce Team'

# ===== HELPERS =====
def set_cell_shading(cell, color_hex):
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color_hex}" w:val="clear"/>')
    cell._tc.get_or_add_tcPr().append(shading)

def set_cell_borders(cell, color='AAAAAA', sz=4, val='single'):
    borders = parse_xml(
        f'<w:tcBorders {nsdecls("w")}>'
        f'  <w:top w:val="{val}" w:sz="{sz}" w:space="0" w:color="{color}"/>'
        f'  <w:bottom w:val="{val}" w:sz="{sz}" w:space="0" w:color="{color}"/>'
        f'  <w:left w:val="{val}" w:sz="{sz}" w:space="0" w:color="{color}"/>'
        f'  <w:right w:val="{val}" w:sz="{sz}" w:space="0" w:color="{color}"/>'
        f'</w:tcBorders>'
    )
    cell._tc.get_or_add_tcPr().append(borders)

def set_paragraph_spacing(p, before=0, after=0):
    pPr = p._p.get_or_add_pPr()
    sp = parse_xml(f'<w:spacing {nsdecls("w")} w:before="{before}" w:after="{after}"/>')
    pPr.append(sp)

def set_table_full_width(table):
    tblW = parse_xml(f'<w:tblW {nsdecls("w")} w:w="5000" w:type="pct"/>')
    table._tbl.tblPr.append(tblW)

def mark_header_row(table):
    trPr = table.rows[0]._tr.get_or_add_trPr()
    trPr.append(parse_xml(f'<w:tblHeader {nsdecls("w")}/>'))

def clean_heading(text):
    t = text.strip()
    t = re.sub(r'^[^\w\s]+\s*', '', t, count=1) if t and not t[0].isalnum() else t
    return t

def add_field(paragraph, field_code):
    run1 = paragraph.add_run()
    run1._r.append(parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="begin"/>'))
    run2 = paragraph.add_run()
    run2._r.append(parse_xml(f'<w:instrText {nsdecls("w")} xml:space="preserve"> {field_code} </w:instrText>'))
    run3 = paragraph.add_run()
    run3._r.append(parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="separate"/>'))
    run4 = paragraph.add_run()
    run5 = paragraph.add_run()
    run5._r.append(parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="end"/>'))
    return run4

def add_rich_runs(paragraph, text, theme, italic=False, color=None):
    pattern = r'(\*\*.*?\*\*|\*[^*]+\*|`[^`]+`)'
    parts = re.split(pattern, text)
    for part in parts:
        if not part:
            continue
        if part.startswith('**') and part.endswith('**'):
            run = paragraph.add_run(part[2:-2])
            run.bold = True
            run.font.name = theme['font_main']
            run.font.size = theme['font_size']
            run.font.color.rgb = color or theme['text']
        elif part.startswith('*') and part.endswith('*') and not part.startswith('**'):
            run = paragraph.add_run(part[1:-1])
            run.italic = True
            run.font.name = theme['font_main']
            run.font.size = theme['font_size']
            run.font.color.rgb = theme['muted']
        elif part.startswith('`') and part.endswith('`'):
            run = paragraph.add_run(part[1:-1])
            run.font.name = theme['font_code']
            run.font.size = theme['code_size']
            run.font.color.rgb = RGBColor(0xC7, 0x25, 0x4E)
        else:
            run = paragraph.add_run(part)
            run.font.name = theme['font_main']
            run.font.size = theme['font_size']
            run.font.color.rgb = color or theme['text']
            if italic:
                run.italic = True

# ===== DOCUMENT SETUP =====
def setup_page(doc):
    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2)
    section.different_first_page_header_footer = True

def setup_styles(doc, theme):
    styles = doc.styles
    normal = styles['Normal']
    normal.font.name = theme['font_main']
    normal.font.size = theme['font_size']
    normal.font.color.rgb = theme['text']
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.15

    heading_cfg = [
        (1, theme['h1_size'], theme['primary']),
        (2, theme['h2_size'], theme['secondary']),
        (3, theme['h3_size'], theme['accent']),
    ]
    for level, size, color in heading_cfg:
        h = styles[f'Heading {level}']
        h.font.name = theme['font_main']
        h.font.size = size
        h.font.color.rgb = color
        h.font.bold = True
        h.paragraph_format.space_before = Pt(18 if level <= 2 else 12)
        h.paragraph_format.space_after = Pt(8)
        if level == 1:
            h.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER

def set_properties(doc, title, author):
    core = doc.core_properties
    core.title = title
    core.author = author
    core.subject = 'User Guide'
    core.keywords = 'ABM, AI, Multi-Agent, Workforce, Antigravity'
    core.language = 'vi-VN'

# ===== COVER PAGE =====
def add_cover_page(doc, theme):
    bar = doc.add_table(rows=1, cols=1)
    bar.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_full_width(bar)
    cell = bar.rows[0].cells[0]
    cell.text = ''
    set_cell_shading(cell, theme['bg_table_header'])
    set_cell_borders(cell, color=theme['bg_table_header'], sz=0, val='none')
    trPr = bar.rows[0]._tr.get_or_add_trPr()
    trPr.append(parse_xml(f'<w:trHeight {nsdecls("w")} w:val="567" w:hRule="exact"/>'))

    for _ in range(4): doc.add_paragraph('')

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('HUONG DAN SU DUNG')
    run.font.size = Pt(36)
    run.font.color.rgb = theme['primary']
    run.font.name = theme['font_main']
    run.bold = True

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('ABM-WORKFORCE v3.0')
    run.font.size = Pt(28)
    run.font.color.rgb = theme['accent']
    run.font.name = theme['font_main']
    run.bold = True

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('=' * 50)
    run.font.color.rgb = theme['accent']
    run.font.size = Pt(12)

    taglines = [
        'Unified Multi-Agent System',
        'BMAD 4-Phase Methodology',
        '10 Jarvis Workers | 140 Skills | 94 Workflows',
    ]
    for line in taglines:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(line)
        run.font.size = Pt(14)
        run.font.color.rgb = theme['muted']
        run.font.name = theme['font_main']
        run.italic = True

    for _ in range(5): doc.add_paragraph('')

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('Cap nhat: Thang 03/2026 - Restructured Workspace Edition')
    run.font.size = Pt(10)
    run.font.color.rgb = theme['muted']

    doc.add_page_break()

# ===== STATIC TABLE OF CONTENTS =====
def scan_headings(lines):
    headings = []
    skip_first_h1 = True
    for line in lines:
        line = line.rstrip()
        if line.startswith('### '):
            headings.append((3, clean_heading(line[4:])))
        elif line.startswith('## '):
            headings.append((2, clean_heading(line[3:])))
        elif line.startswith('# '):
            if skip_first_h1:
                skip_first_h1 = False
                continue
            headings.append((1, clean_heading(line[2:])))
    return headings

def add_static_toc(doc, theme, headings):
    p = doc.add_heading('MUC LUC', level=1)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('=' * 50)
    run.font.color.rgb = theme['accent']
    run.font.size = Pt(8)
    doc.add_paragraph('')

    for level, title in headings:
        p = doc.add_paragraph()
        if level == 2:
            indent, font_size, color, bold = Cm(0), Pt(12), theme['primary'], True
            set_paragraph_spacing(p, before=80, after=40)
        elif level == 3:
            indent, font_size, color, bold = Cm(1.2), Pt(10.5), theme['accent'], False
            set_paragraph_spacing(p, before=20, after=20)
        else:
            indent, font_size, color, bold = Cm(0), Pt(13), theme['primary'], True
            set_paragraph_spacing(p, before=100, after=40)

        p.paragraph_format.left_indent = indent
        if level == 2:
            run = p.add_run('\u25A0  ')
            run.font.size = Pt(8)
            run.font.color.rgb = theme['accent']
        if level == 3:
            run = p.add_run('\u25B8  ')
            run.font.size = Pt(8)
            run.font.color.rgb = theme['muted']

        run = p.add_run(title)
        run.font.name = theme['font_main']
        run.font.size = font_size
        run.font.color.rgb = color
        run.bold = bold

        if level == 2:
            pPr = p._p.get_or_add_pPr()
            pPr.append(parse_xml(
                f'<w:pBdr {nsdecls("w")}>'
                f'  <w:bottom w:val="dotted" w:sz="2" w:space="2" w:color="CCCCCC"/>'
                f'</w:pBdr>'
            ))
    doc.add_page_break()

# ===== HEADERS & FOOTERS =====
def add_headers_footers(doc, theme, title):
    for section in doc.sections:
        section.different_first_page_header_footer = True

        header = section.header
        header.is_linked_to_previous = False
        hp = header.paragraphs[0]
        hp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        run = hp.add_run(title)
        run.font.size = Pt(8)
        run.font.color.rgb = theme['muted']
        run.font.name = theme['font_main']
        run.italic = True
        pPr = hp._p.get_or_add_pPr()
        pPr.append(parse_xml(
            f'<w:pBdr {nsdecls("w")}>'
            f'  <w:bottom w:val="single" w:sz="4" w:space="6" w:color="CCCCCC"/>'
            f'</w:pBdr>'
        ))

        footer = section.footer
        footer.is_linked_to_previous = False
        fp = footer.paragraphs[0]
        fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = fp.add_run('Trang ')
        r.font.size, r.font.color.rgb, r.font.name = Pt(8), theme['muted'], theme['font_main']
        add_field(fp, 'PAGE')
        r = fp.add_run(' / ')
        r.font.size, r.font.color.rgb = Pt(8), theme['muted']
        add_field(fp, 'NUMPAGES')

# ===== PAGE BORDER =====
def add_page_border(doc):
    c = '1B2A4A'
    for section in doc.sections:
        section._sectPr.append(parse_xml(
            f'<w:pgBorders {nsdecls("w")} w:offsetFrom="page">'
            f'  <w:top w:val="single" w:sz="4" w:space="24" w:color="{c}"/>'
            f'  <w:bottom w:val="single" w:sz="4" w:space="24" w:color="{c}"/>'
            f'  <w:left w:val="single" w:sz="4" w:space="24" w:color="{c}"/>'
            f'  <w:right w:val="single" w:sz="4" w:space="24" w:color="{c}"/>'
            f'</w:pgBorders>'
        ))

# ===== CONTENT ELEMENTS =====
def add_styled_table(doc, theme, headers, rows):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True
    set_table_full_width(table)

    for idx, ht in enumerate(headers):
        cell = table.rows[0].cells[idx]
        cell.text = ''
        p = cell.paragraphs[0]
        run = p.add_run(re.sub(r'\*\*([^*]+)\*\*', r'\1', ht.strip()))
        run.font.name, run.font.size, run.font.bold, run.font.color.rgb = theme['font_main'], Pt(10), True, theme['white']
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        set_cell_shading(cell, theme['bg_table_header'])
        set_cell_borders(cell, color=theme['bg_table_header'], sz=6)

    for ri, row_data in enumerate(rows):
        for ci, ct in enumerate(row_data):
            cell = table.rows[ri + 1].cells[ci]
            cell.text = ''
            add_rich_runs(cell.paragraphs[0], ct.strip(), theme)
            if ri % 2 == 1: set_cell_shading(cell, theme['bg_table_alt'])
            set_cell_borders(cell, color=theme['border_light'], sz=4)

    mark_header_row(table)
    doc.add_paragraph('')

def add_code_block(doc, theme, code_text):
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_full_width(table)
    cell = table.rows[0].cells[0]
    cell.text = ''
    set_cell_shading(cell, theme['bg_code'])
    set_cell_borders(cell, color=theme['border_light'], sz=4)

    for line in code_text.split('\n'):
        p = cell.add_paragraph()
        run = p.add_run(line)
        run.font.name, run.font.size, run.font.color.rgb = theme['font_code'], theme['code_size'], theme['text']
        set_paragraph_spacing(p, before=0, after=0)

    if cell.paragraphs and cell.paragraphs[0].text == '':
        cell.paragraphs[0]._p.getparent().remove(cell.paragraphs[0]._p)
    doc.add_paragraph('')

def add_blockquote(doc, theme, text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(1.2)
    pPr = p._p.get_or_add_pPr()
    pPr.append(parse_xml(
        f'<w:pBdr {nsdecls("w")}><w:left w:val="single" w:sz="18" w:space="12" w:color="3498DB"/></w:pBdr>'
    ))
    add_rich_runs(p, text, theme, italic=True, color=theme['muted'])

def add_divider(doc, theme):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_spacing(p, before=100, after=100)
    pPr = p._p.get_or_add_pPr()
    pPr.append(parse_xml(
        f'<w:pBdr {nsdecls("w")}><w:bottom w:val="single" w:sz="6" w:space="1" w:color="CCCCCC"/></w:pBdr>'
    ))
    run = p.add_run(' ')
    run.font.size = Pt(2)

def add_rich_paragraph(doc, theme, text, is_bullet=False, is_numbered=False):
    p = doc.add_paragraph()
    if is_bullet: p.style = doc.styles['List Bullet']
    elif is_numbered: p.style = doc.styles['List Number']
    add_rich_runs(p, text, theme)
    return p

# ===== MAIN PARSER =====
def parse_and_build(md_path, output_path):
    with open(md_path, 'r', encoding='utf-8') as f:
        lines = f.readlines()

    headings = scan_headings(lines)
    doc = Document()
    setup_page(doc)
    setup_styles(doc, THEME)
    set_properties(doc, DOC_TITLE, DOC_AUTHOR)
    add_cover_page(doc, THEME)
    add_static_toc(doc, THEME, headings)

    i = 0
    first_h1 = True
    h2_count = 0

    while i < len(lines):
        line = lines[i].rstrip('\n').rstrip('\r')

        if line.startswith('# ') and first_h1:
            first_h1 = False
            i += 1
            continue
        if line.strip() == '---':
            add_divider(doc, THEME)
            i += 1
            continue
        if line.startswith('### '):
            doc.add_heading(clean_heading(line[4:]), level=3)
            i += 1
            continue
        if line.startswith('## '):
            h2_count += 1
            if h2_count > 1: doc.add_page_break()
            doc.add_heading(clean_heading(line[3:]), level=2)
            i += 1
            continue
        if line.startswith('# '):
            doc.add_heading(clean_heading(line[2:]), level=1)
            i += 1
            continue
        if line.strip().startswith('```'):
            code_lines = []
            i += 1
            while i < len(lines) and not lines[i].strip().startswith('```'):
                code_lines.append(lines[i].rstrip('\n').rstrip('\r'))
                i += 1
            i += 1
            add_code_block(doc, THEME, '\n'.join(code_lines))
            continue
        if '|' in line and i + 1 < len(lines) and '---' in lines[i + 1]:
            headers = [c.strip() for c in line.split('|') if c.strip()]
            i += 2
            rows = []
            while i < len(lines) and '|' in lines[i] and lines[i].strip():
                row = [c.strip() for c in lines[i].split('|') if c.strip()]
                while len(row) < len(headers): row.append('')
                rows.append(row[:len(headers)])
                i += 1
            add_styled_table(doc, THEME, headers, rows)
            continue
        if line.startswith('>'):
            text = line.lstrip('>').strip()
            if text: add_blockquote(doc, THEME, text)
            i += 1
            continue
        if line.startswith('- '):
            add_rich_paragraph(doc, THEME, line[2:].strip(), is_bullet=True)
            i += 1
            continue
        m = re.match(r'^(\d+)\.\s+(.*)', line)
        if m:
            add_rich_paragraph(doc, THEME, m.group(2).strip(), is_numbered=True)
            i += 1
            continue
        if not line.strip():
            i += 1
            continue

        add_rich_paragraph(doc, THEME, line)
        i += 1

    add_headers_footers(doc, THEME, DOC_TITLE)
    add_page_border(doc)

    doc.save(output_path)
    print(f'OK: {output_path}')

if __name__ == '__main__':
    md = sys.argv[1] if len(sys.argv) > 1 else 'D:/AntigravityWorkspace/ABM-Workforce-HDSD.md'
    out = sys.argv[2] if len(sys.argv) > 2 else 'D:/AntigravityWorkspace/ABM-Workforce-HDSD-v4.docx'
    parse_and_build(md, out)
